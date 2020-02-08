from psytest_tools import b64dec, decrypt, get_testees_by_grade_not_yet, make_filename, get_testees_by_grade_done
from application import app
from application import decorators as decors
from flask import session, render_template, url_for, redirect, send_file
from os import path
import docx




@app.route('/admin_download/<name>/<target>')
@decors.check_admin
def admin_download(name, target):
    dec_name = b64dec(name)

    if target == 'not_yet':
        testees = tuple(get_testees_by_grade_not_yet(session['psy_login'], dec_name))
        doc = docx.Document()
        table = doc.add_table(rows=len(testees), cols=3)
        table.style = 'Table Grid'
        for row, testee in enumerate(testees):
            table.cell(row, 0).text = testee['login']
            table.cell(row, 1).text = decrypt(testee['pas'])
            table.cell(row, 2).text = '\n\n'

    if target == 'done':
        testees = tuple(get_testees_by_grade_done(session['psy_login'], dec_name))
        doc = docx.Document()
        table = doc.add_table(rows=len(testees), cols=3)
        table.style = 'Table Grid'
        for row, testee in enumerate(testees):
            table.cell(row, 0).text = testee['login']
            table.cell(row, 1).text = testee['result']
            table.cell(row, 2).text = '\n\n'

    filename = path.join(app.config['DOCKS_FOLDER'], make_filename(session['psy_login'], 'docx'))
    doc.save(filename)
    return send_file(filename, cache_timeout=0, as_attachment=True)